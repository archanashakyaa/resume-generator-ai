from flask import Flask, request, send_file, jsonify, send_from_directory, render_template
from flask_cors import CORS
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.colors import HexColor
import os
import traceback
import logging
import time
import uuid
import re
import json
from dotenv import load_dotenv
import requests
from bs4 import BeautifulSoup

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# CORS configuration
CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"]
    }
})

# Groq API Configuration
GROQ_API_KEY = os.getenv('GROQ_API_KEY', '')
GROQ_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

# Initialize Groq client
client = None
if GROQ_API_KEY:
    try:
        client = Groq(api_key=GROQ_API_KEY)
        test_response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": "test"}],
            max_tokens=5
        )
        logger.info(f"Groq API connected successfully with model: {GROQ_MODEL}")
    except Exception as e:
        logger.error(f"Groq API connection failed: {e}")
        client = None
else:
    logger.error("No GROQ_API_KEY found")


# Resume Enhancement Prompts
GLOBAL_RULES = [
    "Use a professional, employer-focused tone.",
    "Do not use first-person pronouns.",
    "Be concise, quantifiable, and clear.",
    "Fix grammar, avoid redundancy.",
    "Do not invent experiences or education.",
]

GLOBAL_RULE = "Global Resume Rules:\n" + "\n".join(f"{i + 1}. {rule}" for i, rule in enumerate(GLOBAL_RULES)) + "\n"

resume_prompts = {
    "summary": (
        "Role: Expert Resume Consultant & AI Career Coach Assistant\n"
        "Objective: Generate a concise, impactful, and professionally written Resume Summary tailored to the user's role, skills, and career focus.\n\n"

        "--------------------\n"
        "🎯 **PRIMARY TASK:**\n"
        "Craft a polished professional summary based solely on the user's provided information. The tone should be confident, clear, and employer-focused.\n"
        "Do not invent experiences, education, or achievements beyond what is stated.\n\n"

        "--------------------\n"
        "📋 **CONTENT GUIDELINES:**\n"
        "1. Rewrite the summary in **50–70 words (2–4 concise sentences)**.\n"
        "2. Begin with the **professional title** and include **years of relevant experience** only if explicitly provided.\n"
        "3. Emphasize **key technical and professional skills**, **core strengths**, and **career value**.\n"
        "4. Include 1–2 **measurable or outcome-driven achievements** if contextually appropriate.\n"
        "5. Use a variety of strong professional adjectives such as *results-driven, strategic, analytical, innovative, adaptable, or detail-oriented*.\n"
        "6. Ensure the writing style aligns with **ATS optimization** and targeted job relevance.\n"
        "7. Maintain grammatical precision, avoid redundancy, and exclude first-person language.\n"
        "8. Always return **only one final, polished summary paragraph** — no explanations, introductions, or commentary.\n\n"

        "--------------------\n"
        "🧩 **KEY EXTRACTION RULES:**\n"
        "From the user's input, identify and extract the following details:\n"
        "- **Job Role / Target Title** (e.g., Project Manager, Data Analyst)\n"
        "- **Core Skills / Technologies** (e.g., Python, SQL, Agile, CRM)\n"
        "- **Years of Experience** (only if explicitly mentioned; do not infer or assume)\n\n"

        "--------------------\n"
        "🧠 **SUMMARY GENERATION LOGIC:**\n"
        "- **If years of experience are mentioned:** naturally incorporate them into the first sentence.\n"
        "  *Example:* 'Results-driven Project Manager with **6 years of experience** in delivering cross-functional initiatives...'\n"
        "- **If no experience is mentioned:** use experience-neutral phrasing.\n"
        "  *Use openings such as:* 'Proficient in...', 'Skilled in...', 'Adept at...', 'A motivated [Job Role] with expertise in...', "
        "'Demonstrates strong capabilities in...'\n"
        "- Avoid repetitive openings like 'A highly motivated...'\n\n"

        "--------------------\n"
        "🧾 **OUTPUT FORMAT:**\n"
        "- Deliver only the final professional summary paragraph.\n"
        "- The output must be 3–4 sentences, formatted as a single paragraph.\n"
        "- Exclude any system messages, acknowledgments, or preambles.\n\n"

        "--------------------\n"
        "💡 **EXAMPLES:**\n"
        "**Example 1 (With Experience):**\n"
        "User Input: 'Write a summary for a Senior Software Engineer with 8 years of experience in Java and cloud technologies.'\n"
        "Output: 'Seasoned Senior Software Engineer with over 8 years of expertise in developing robust applications using Java and cloud-native frameworks. "
        "Recognized for leading cross-functional teams to deliver scalable, high-performance systems. Adept at optimizing full-stack architectures and implementing "
        "best practices in DevOps and cloud deployment.'\n\n"

        "**Example 2 (Without Experience):**\n"
        "User Input: 'I need a professional summary for a Marketing Coordinator skilled in social media, content creation, and SEO.'\n"
        "Output: 'Creative and data-driven Marketing Coordinator skilled in crafting engaging content and managing multi-platform campaigns. "
        "Proficient in leveraging SEO and analytics to enhance digital presence and audience engagement. Demonstrates strong adaptability and strategic thinking to "
        "align marketing initiatives with brand growth goals.'\n\n"

        "--------------------\n"
        "**Now, generate a concise, professional summary based on the user's query below:**"
    ),

    "experience": (
"Role: Expert Resume Consultant & Career Strategist\n"
    "Objective: Transform raw work experience into compelling, achievement-oriented narratives that demonstrate measurable impact, transferable skills, and clear value to recruiters and ATS systems.\n\n"

    "**PROCESSING FRAMEWORK:**\n\n"

    "1.  **Input Analysis & Gap Handling:**\n"
    "    - Extract role title, company, timeframe, and initial responsibilities.\n"
    "    - If dates/timeframes are missing, structure chronologically based on context.\n"
    "    - If specific metrics are absent, infer reasonable scope based on role seniority and industry norms.\n\n"

    "2.  **Content Transformation Rules:**\n"
    "    - Convert each role into **70–120 word, achievement-focused narratives**.\n"
    "    - Structure with **3–5 concise, impactful bullet points** per role.\n"
    "    - Begin each bullet with **strong, varied action verbs** (e.g., Managed, Led, Optimized, Implemented, Streamlined, Spearheaded, Orchestrated, Pioneered).\n"
    "    - Apply **CAR (Challenge-Action-Result)** or **STAR (Situation-Task-Action-Result)** methodology implicitly.\n\n"

    "3.  **Quantification Protocol:**\n"
    "    - **Explicit Metrics:** Use provided numbers directly (revenue, percentages, timeframes).\n"
    "    - **Inferred Impact:** Where numbers are missing, use industry-appropriate scale indicators (e.g., 'large-scale', 'multiple', 'cross-functional').\n"
    "    - **Scope Indicators:** Quantify team sizes, budgets, project scales, or client impact based on role level.\n\n"

    "4.  **ATS & Readability Optimization:**\n"
    "    - Incorporate industry-specific keywords and transferable skills aligned with target job description.\n"
    "    - Use plain-text, ATS-friendly bullets only (dash - or •).\n"
    "    - Maintain a professional tone without personal pronouns or opinions.\n"
    "    - Ensure grammatical consistency (past tense for prior roles, present tense for current roles).\n\n"

    "5.  **Strategic Positioning:**\n"
    "    - Align achievements with common resume screening criteria.\n"
    "    - Highlight leadership, problem-solving, technical competencies, and business impact.\n"
    "    - Emphasize measurable results and ROI whenever possible.\n\n"

    "6.  **Output Standards:**\n"
    "    - Return a **single polished version per role**.\n"
    "    - Maintain consistent formatting across all experiences.\n"
    "    - Ensure clarity, professionalism, and factual accuracy.\n"
    "    - Follow all **Global Resume Rules**.\n\n"

    "**FORMAT TEMPLATE:**\n"
    "- [Achievement 1: Action verb + quantified result + business impact]\n"
    "- [Achievement 2: Leadership/initiative + scope + outcome]\n"
    "- [Achievement 3: Process improvement + metrics + efficiency gain]\n\n"

    "**EXAMPLE TRANSFORMATIONS:**\n\n"
    "*Before:*\n"
    "\"Was responsible for sales team and meeting targets\"\n\n"
    "*After:*\n"
    "- Led and motivated 12-person sales team to exceed quarterly targets by 15-25% for 6 consecutive quarters\n"
    "- Implemented a new CRM system that improved sales pipeline visibility and increased conversion rate by 30%\n"
    "- Developed strategic account plans that expanded key client portfolios by 40% year-over-year\n\n"
    "*Before:*\n"
    "\"Handled software development and project management\"\n\n"
    "*After:*\n"
    "- Spearheaded development of 3 major product features using Agile methodology, delivering 2 weeks ahead of schedule\n"
    "- Managed $500K project budget while coordinating cross-functional teams of 15+ engineers and designers\n"
    "- Optimized deployment processes, reducing production incidents by 60% and improving system reliability\n\n"

    "**READY TO PROCESS:**\n"
    "[User's work experience input will be transformed here]"
    ),
    "skills": (
        "Role: Expert Resume Consultant\n"
        "Objective: Refine the Skills section for ATS optimization and employer appeal.\n"
        "Instruction 1: Rewrite as a concise list, limited to 10–15 core skills, using comma-separated format or grouped categories (e.g., Technical Skills, Soft Skills, Certifications).\n"
        "Instruction 2: Prioritize hard skills, technical competencies, certifications, and industry-specific terminology mentioned in the job description.\n"
        "Instruction 3: Use exact keywords and phrases from the posting, including both spelled-out and acronym forms (e.g., Search Engine Optimization (SEO)).\n"
        "Instruction 4: Avoid vague or outdated terms; keep all skills current, specific, and relevant to the targeted job.\n"
        "Instruction 5: Eliminate redundancy and ensure logical grouping for easy readability.\n"
        "Instruction 6: Integrate the most critical skills naturally in both this section and throughout work experience, avoiding keyword stuffing.\n"
        "Instruction 7: Ensure formatting is consistent, professional, and ATS-friendly.\n"
        "Instruction 8: Return only the single best version for the resume.\n"
        "Notes: Follow all Global Resume Rules."
    ),
    "education": (
        "Role: Expert Resume Consultant\n"
        "Objective: Improve the Education section to clearly showcase qualifications relevant to the targeted role.\n"
        "Instruction 1: Rewrite in 50–100 words.\n"
        "Instruction 2: Clearly present degrees, certifications, relevant coursework, honors, and any notable academic achievements.\n"
        "Instruction 3: Highlight how the education supports career goals and aligns with the targeted job.\n"
        "Instruction 4: Use concise, professional language; maintain clarity, proper grammar, and factual accuracy.\n"
        "Instruction 5: Ensure formatting is ATS-friendly and easily scannable.\n"
        "Instruction 6: Return only the single best, polished version for the resume.\n"
        "Notes: Follow all Global Resume Rules."
    ),
    "projects": (
        "Role: Expert Resume Consultant & Project Portfolio Strategist\n"
        "Objective: Transform the Projects section to maximize clarity, relevance, and measurable impact.\n"
        "Instructions:\n"
        "1. For EACH project, enhance the title and description separately, ensuring clarity and professionalism.\n"
        "2. Provide a concise 2–3 sentence description that communicates the project's purpose, scope, and relevance to the targeted role.\n"
        "3. Clearly list key skills, technologies, and tools applied in the project.\n"
        "4. Emphasize specific tasks, responsibilities, and contributions using strong action verbs.\n"
        "5. Quantify achievements wherever possible to demonstrate measurable impact (e.g., improved performance by 20%, reduced processing time by 30%).\n"
        "6. Maintain concise phrasing, readability, and professional formatting.\n"
        "7. Tailor language to align with the target job description and ATS-relevant keywords.\n"
        "8. Ensure flawless grammar, spelling, and consistent formatting throughout.\n"
        "9. Format output EXACTLY as:\n"
        "   Title: [Enhanced Project Title]\n"
        "   Description: [Enhanced description in 2-3 sentences]\n"
        "   ---\n"
        "10. Return ONLY the enhanced project content; do NOT include explanations, commentary, or extra text.\n"
        "Notes: Follow all Global Resume Rules and best practices for professional project presentation."
    )
}


def sanitize_input(text, max_chars=3000):
    """Clean and limit input text."""
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text).strip()
    if len(text) > max_chars:
        text = text[:max_chars].rsplit(' ', 1)[0]
    return text


def clean_ai_response(text):
    """Remove common AI response artifacts."""
    if not text:
        return ""

    # Remove code fences
    text = re.sub(r'^```(?:\w+)?\s*|```$', '', text, flags=re.MULTILINE).strip()

    # Remove common preambles
    patterns = [
        r'^(?:Here\'s|Here is|Enhanced version:|Enhanced:|Sure,?.*?:)\s*',
        r'^(?:Certainly|Of course|Absolutely).*?:\s*',
        r'^\*\*.*?\*\*\s*',  # Remove markdown bold headers
    ]
    for pattern in patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)

    # Remove quotes
    text = re.sub(r'^["\']|["\']$', '', text.strip())

    return text.strip()


def enhance_section(section_name, content, max_retries=2):
    """Enhance a resume section using Groq AI with your specific prompts."""
    if not client:
        logger.error("Groq client not available")
        return content

    section_name = section_name.lower().strip()

    # Handle projects - parse JSON if provided
    if section_name == "projects":
        try:
            projects = json.loads(content)
            if isinstance(projects, list) and projects:
                formatted = "\n\n".join([
                    f"Project {i + 1}:\nTitle: {p.get('title', 'Untitled')}\n"
                    f"Description: {p.get('description', 'No description')}"
                    for i, p in enumerate(projects)
                ])
                content = formatted
        except (json.JSONDecodeError, TypeError):
            pass

    # Sanitize input
    content = sanitize_input(content)
    if not content:
        logger.warning(f"Empty content for section: {section_name}")
        return ""

    # Get the detailed prompt
    prompt_template = resume_prompts.get(section_name, resume_prompts["summary"])

    # Construct full prompt with global rules
    full_prompt = (
        f"{GLOBAL_RULE}\n\n"
        f"{prompt_template}\n\n"
        f"User Input:\n{content}\n\n"
        f"Enhanced Content:"
    )

    # Retry logic with exponential backoff
    for attempt in range(max_retries + 1):
        try:
            logger.info(f"Enhancing {section_name} (attempt {attempt + 1}/{max_retries + 1})")

            response = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert resume consultant. Follow the instructions precisely and return ONLY the enhanced content without any preambles, explanations, or meta-commentary."
                    },
                    {
                        "role": "user",
                        "content": full_prompt
                    }
                ],
                temperature=0.5,
                max_tokens=1024,
                top_p=0.95
            )

            enhanced = response.choices[0].message.content.strip()
            enhanced = clean_ai_response(enhanced)

            if not enhanced:
                raise ValueError("Empty response from AI")

            logger.info(f"Successfully enhanced {section_name} ({len(enhanced)} chars)")
            return enhanced

        except Exception as e:
            logger.error(f"Enhancement failed (attempt {attempt + 1}): {str(e)}")
            if attempt >= max_retries:
                logger.warning(f"Max retries reached, returning original content for {section_name}")
                return content
            time.sleep(1 * (2 ** attempt))

    return content


def format_for_docx(text):
    """Format text into paragraphs for DOCX."""
    if not text:
        return

    text = text.strip()

    # Handle comma-separated lists (skills)
    if ',' in text and '\n' not in text and len(text.split(',')) > 2:
        yield text
        return

    # Split by double newlines or "---"
    for block in re.split(r'\n\s*\n|---', text):
        block = block.strip()
        if not block:
            continue

        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]

        # Check if it's a bullet list
        is_list = all(re.match(r'^[•\-]\s+', ln) for ln in lines if ln)

        if is_list and len(lines) > 1:
            yield '\n'.join(lines)
        else:
            yield ' '.join(lines)


def create_enhanced_docx(resume_data, filename=None):
    """Create a professionally formatted DOCX resume."""
    if not filename:
        filename = f"Resume_{uuid.uuid4().hex[:8]}.docx"

    os.makedirs("generated", exist_ok=True)
    filepath = os.path.join("generated", filename)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add name as title
    name = resume_data.get('Name', '').strip()
    if name:
        title = doc.add_heading(name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section order
    section_order = [
        'Contact Information',
        'Professional Summary',
        'Work Experience',
        'Education',
        'Skills',
        'Projects'
    ]

    # Add sections
    for section_name in section_order:
        content = resume_data.get(section_name, '').strip()
        if not content or content.startswith('['):
            continue

        # Add section heading
        heading = doc.add_heading(section_name, level=1)
        heading_format = heading.runs[0].font
        heading_format.color.rgb = RGBColor(31, 78, 121)

        # Add content
        for paragraph_text in format_for_docx(content):
            para = doc.add_paragraph(paragraph_text)
            para.paragraph_format.space_after = Pt(6)

    doc.save(filepath)
    logger.info(f"DOCX saved: {filepath}")
    return filepath


def create_enhanced_pdf(resume_data, filename=None):
    """Create a professionally formatted PDF resume."""
    if not filename:
        filename = f"Resume_{uuid.uuid4().hex[:8]}.pdf"

    os.makedirs("generated", exist_ok=True)
    filepath = os.path.join("generated", filename)

    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            topMargin=0.5 * inch, bottomMargin=0.5 * inch,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch)

    # Custom styles
    styles = getSampleStyleSheet()

    # Name style
    name_style = ParagraphStyle(
        'CustomName',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=HexColor('#1F4E79'),
        alignment=TA_CENTER,
        spaceAfter=6
    )

    # Contact style
    contact_style = ParagraphStyle(
        'ContactInfo',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=12
    )

    # Section heading style
    section_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=HexColor('#1F4E79'),
        spaceBefore=12,
        spaceAfter=6,
        borderWidth=0,
        borderColor=HexColor('#1F4E79'),
        borderPadding=0
    )

    # Body text style
    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=6
    )

    # Build document
    story = []

    # Add name
    name = resume_data.get('Name', '').strip()
    if name:
        story.append(Paragraph(name, name_style))

    # Add contact information
    contact = resume_data.get('Contact Information', '').strip()
    if contact:
        story.append(Paragraph(contact, contact_style))

    story.append(Spacer(1, 0.1 * inch))

    # Section order
    section_order = [
        'Professional Summary',
        'Work Experience',
        'Education',
        'Skills',
        'Projects'
    ]

    # Add sections
    for section_name in section_order:
        content = resume_data.get(section_name, '').strip()
        if not content or content.startswith('['):
            continue

        # Add section heading
        story.append(Paragraph(section_name, section_style))

        # Add content
        # Handle bullet points
        content_lines = content.split('\n')
        for line in content_lines:
            line = line.strip()
            if not line:
                continue

            # Convert bullet points
            if line.startswith('•') or line.startswith('-'):
                line = '&bull; ' + line[1:].strip()

            story.append(Paragraph(line, body_style))

        story.append(Spacer(1, 0.1 * inch))

    # Build PDF
    doc.build(story)
    logger.info(f"PDF saved: {filepath}")
    return filepath


# -----------------------------------------------------------------------------
# Route: /
# Method: GET
# Purpose: Serve the main resume builder page
# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------
# Route: /
# Method: GET
# Purpose: Serve the main resume builder page
# -----------------------------------------------------------------------------
@app.route("/")
def index():
    """Serve the main resume builder page."""
    return render_template('index.html')


# -----------------------------------------------------------------------------
# Route: /health
# Method: GET
# Purpose: Health check endpoint to verify server is running
# -----------------------------------------------------------------------------
@app.route("/health", methods=["GET"])
def health():
    """Health check endpoint."""
    return jsonify({
        "status": "ok",
        "groq_connected": client is not None,
        "model": GROQ_MODEL
    }), 200


# -----------------------------------------------------------------------------
# Route: /enhance
# Method: POST
# Purpose: Enhance a specific section of the resume using AI
# -----------------------------------------------------------------------------
@app.route("/enhance", methods=["POST"])
def enhance():
    """Enhance a specific resume section."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        section = data.get('section', '').strip()
        content = data.get('content', '').strip()

        if not section or not content:
            return jsonify({"error": "Missing section or content"}), 400

        logger.info(f"Enhancing section: {section}")
        enhanced = enhance_section(section, content)

        return jsonify({
            "success": True,
            "section": section,
            "enhanced": enhanced
        }), 200

    except Exception as e:
        logger.error(f"Enhancement error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# -----------------------------------------------------------------------------
# Route: /generate_resume
# Method: POST
# Purpose: Generate a complete resume with AI enhancements
# -----------------------------------------------------------------------------
@app.route("/generate_resume", methods=["POST"])
def generate_resume():
    """Generate enhanced resume in both DOCX and PDF formats."""
    try:
        resume_data = request.get_json()
        if not resume_data:
            return jsonify({"error": "No resume data provided"}), 400

        logger.info("Generating enhanced resume")

        # Create both formats
        docx_path = create_enhanced_docx(resume_data)
        pdf_path = create_enhanced_pdf(resume_data)

        return jsonify({
            "success": True,
            "message": "Resume generated successfully",
            "docx": os.path.basename(docx_path),
            "pdf": os.path.basename(pdf_path)
        }), 200

    except Exception as e:
        logger.error(f"Resume generation error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# -----------------------------------------------------------------------------
# Route: /download
# Method: GET
# Purpose: Download the most recently generated resume in DOCX format.
# -----------------------------------------------------------------------------
@app.route("/download", methods=["GET"])
def download():
    """Download the most recent resume in DOCX format."""
    try:
        # Check if the 'generated' directory exists
        if not os.path.exists('generated'):
            # If not, return a 404 error with a message
            return jsonify({"error": "No resumes generated yet"}), 404

        # List all files in 'generated' that end with .docx (Word files)
        files = [f for f in os.listdir('generated') if f.endswith('.docx')]
        if not files:
            # If no DOCX files found, return a 404 error
            return jsonify({"error": "No resume found"}), 404

        # Find the most recently created DOCX file (by creation time)
        latest = max([os.path.join('generated', f) for f in files], key=os.path.getctime)
        # Send the file as a download with a user-friendly filename
        return send_file(latest, as_attachment=True, download_name='Enhanced_Resume.docx')

    except Exception as e:
        # Log any errors and return a 500 error with the error message
        logger.error(f"Download error: {str(e)}")
        return jsonify({"error": str(e)}), 500


# -----------------------------------------------------------------------------
# Route: /download_pdf
# Method: GET
# Purpose: Download the most recently generated resume in PDF format.
# -----------------------------------------------------------------------------
@app.route("/download_pdf", methods=["GET"])
def download_pdf():
    """Download the most recent resume in PDF format."""
    try:
        # Check if the 'generated' directory exists
        if not os.path.exists('generated'):
            return jsonify({"error": "No resumes generated yet"}), 404

        # List all files in 'generated' that end with .pdf
        files = [f for f in os.listdir('generated') if f.endswith('.pdf')]
        if not files:
            return jsonify({"error": "No PDF resume found"}), 404

        # Find the most recently created PDF file
        latest = max([os.path.join('generated', f) for f in files], key=os.path.getctime)
        # Send the file as a download with a user-friendly filename
        return send_file(latest, as_attachment=True, download_name='Enhanced_Resume.pdf')

    except Exception as e:
        logger.error(f"Download PDF error: {str(e)}")
        return jsonify({"error": str(e)}), 500


# -----------------------------------------------------------------------------
# Route: /import_linkedin
# Method: POST
# Purpose: Import profile data from LinkedIn URL or public profile
# -----------------------------------------------------------------------------
@app.route("/import_linkedin", methods=["POST"])
def import_linkedin():
    """Import LinkedIn profile data (simplified version using public data)."""
    try:
        data = request.json
        linkedin_url = data.get("linkedin_url", "")
        
        if not linkedin_url:
            return jsonify({"error": "LinkedIn URL is required"}), 400
        
        # Note: This is a simplified implementation
        # For production, you'd use LinkedIn API with OAuth
        # For now, we'll use AI to extract data from user-provided info
        
        profile_data = {
            "name": "",
            "headline": "",
            "summary": "",
            "experience": [],
            "education": [],
            "skills": [],
            "imported": True,
            "message": "LinkedIn integration ready. Please manually enter your information or use the AI enhancement feature."
        }
        
        return jsonify(profile_data), 200
        
    except Exception as e:
        logger.error(f"LinkedIn import error: {str(e)}")
        return jsonify({"error": str(e)}), 500


# -----------------------------------------------------------------------------
# Route: /analyze_resume
# Method: POST
# Purpose: Analyze resume content and provide scoring with recommendations
# -----------------------------------------------------------------------------
@app.route("/analyze_resume", methods=["POST"])
def analyze_resume():
    """Analyze resume and provide comprehensive scoring with recommendations."""
    try:
        data = request.json
        
        # Extract resume data
        personal_info = data.get("personalInfo", {})
        experiences = data.get("experience", [])
        education = data.get("education", [])
        skills = data.get("skills", [])
        
        # Prepare analysis prompt
        resume_content = f"""
Personal Information:
- Name: {personal_info.get('fullName', 'Not provided')}
- Email: {personal_info.get('email', 'Not provided')}
- Phone: {personal_info.get('phone', 'Not provided')}
- LinkedIn: {personal_info.get('linkedin', 'Not provided')}
- Summary: {personal_info.get('summary', 'Not provided')}

Work Experience ({len(experiences)} entries):
{chr(10).join([f"- {exp.get('title', '')} at {exp.get('company', '')} ({exp.get('startDate', '')} - {exp.get('endDate', '')})" for exp in experiences])}

Education ({len(education)} entries):
{chr(10).join([f"- {edu.get('degree', '')} in {edu.get('field', '')} from {edu.get('school', '')}" for edu in education])}

Skills ({len(skills)} skills):
{', '.join(skills)}
"""

        analysis_prompt = f"""You are an expert resume reviewer and career coach. Analyze the following resume and provide a comprehensive scoring and recommendations.

{resume_content}

Provide your analysis in the following JSON format:
{{
    "overallScore": <number 0-100>,
    "scores": {{
        "formatting": <number 0-100>,
        "content": <number 0-100>,
        "keywords": <number 0-100>,
        "impact": <number 0-100>,
        "completeness": <number 0-100>
    }},
    "strengths": [<list of 3-5 key strengths>],
    "improvements": [<list of 3-5 specific improvements needed>],
    "atsCompatibility": <number 0-100>,
    "recommendations": {{
        "critical": [<list of critical issues to fix>],
        "suggested": [<list of suggested enhancements>],
        "keywords": [<list of missing important keywords>]
    }},
    "summary": "<brief 2-3 sentence summary of the resume quality>"
}}

Be specific, actionable, and constructive in your feedback. Focus on ATS optimization, quantifiable achievements, and professional presentation.
"""

        if not client:
            # Fallback scoring when AI is not available
            return jsonify({
                "overallScore": 75,
                "scores": {
                    "formatting": 80,
                    "content": 70,
                    "keywords": 75,
                    "impact": 70,
                    "completeness": 80
                },
                "strengths": [
                    "Clear contact information provided",
                    f"{len(experiences)} work experience entries included",
                    f"{len(skills)} skills listed"
                ],
                "improvements": [
                    "Add AI-enhanced content for better impact",
                    "Include more quantifiable achievements",
                    "Optimize for ATS compatibility"
                ],
                "atsCompatibility": 75,
                "recommendations": {
                    "critical": ["Use AI enhancement to improve content quality"],
                    "suggested": ["Add more specific metrics and achievements", "Include relevant industry keywords"],
                    "keywords": ["leadership", "project management", "data analysis"]
                },
                "summary": "Good foundation with room for improvement. Use AI enhancement features to optimize content."
            }), 200

        # Get AI analysis
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert resume reviewer. Provide detailed, actionable feedback in valid JSON format only."},
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )

        ai_response = response.choices[0].message.content.strip()
        
        # Extract JSON from response
        json_match = re.search(r'\{[\s\S]*\}', ai_response)
        if json_match:
            analysis_result = json.loads(json_match.group())
        else:
            # Fallback if JSON parsing fails
            analysis_result = {
                "overallScore": 75,
                "scores": {
                    "formatting": 80,
                    "content": 75,
                    "keywords": 70,
                    "impact": 75,
                    "completeness": 80
                },
                "strengths": ["Professional structure", "Complete sections", "Clear formatting"],
                "improvements": ["Add more quantifiable results", "Enhance with action verbs", "Include relevant keywords"],
                "atsCompatibility": 75,
                "recommendations": {
                    "critical": ["Optimize for ATS scanning"],
                    "suggested": ["Add metrics to achievements", "Use industry-specific terminology"],
                    "keywords": ["leadership", "innovation", "results-driven"]
                },
                "summary": "Solid resume with potential for enhancement through AI optimization."
            }
        
        logger.info(f"Resume analyzed with overall score: {analysis_result.get('overallScore', 'N/A')}")
        return jsonify(analysis_result), 200

    except Exception as e:
        logger.error(f"Resume analysis error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# -----------------------------------------------------------------------------
# Main entry point for running the Flask app
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    PORT = 5002  # Changed to avoid conflicts
    
    # Print server startup info to the console
    print("=" * 70)
    print("  Resume Builder Backend Server")
    print("=" * 70)
    print(f"  Model: {GROQ_MODEL}")
    print(f"  API Key: {'Configured' if GROQ_API_KEY else 'Missing'}")
    print(f"  Groq Client: {'Connected' if client else 'Failed'}")
    print(f"  Server: http://localhost:{PORT}")
    print(f"  Health Check: http://localhost:{PORT}/health")
    print("=" * 70)

    # Start the Flask development server, accessible on all network interfaces
    app.run(debug=True, host='0.0.0.0', port=PORT)